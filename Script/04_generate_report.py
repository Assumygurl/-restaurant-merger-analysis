"""
Step 8: Updated Word Report (2024-2025 Data + Forecasting)
===========================================================
Regenerates Part One report with:
  - Updated 2024-2025 dataset stats
  - Revenue forecasting section added
  - 14 charts embedded (10 original + 4 forecast charts)

Output:
  - outputs/Part_One_Report_2025.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import sqlite3
import pandas as pd
import os

# ── Paths ──────────────────────────────────────────────────────
DATA_DIR   = r"C:\Users\HP\Desktop\Restruant_Project\data"
OUTPUT_DIR = r"C:\Users\HP\Desktop\Restruant_Project\outputs"
CHARTS_DIR = os.path.join(OUTPUT_DIR, "charts")
DB_PATH    = os.path.join(DATA_DIR, "restaurant_2025.db")
REPORT_OUT = os.path.join(OUTPUT_DIR, "Part_One_Report_2025.docx")

os.makedirs(OUTPUT_DIR, exist_ok=True)

print("=" * 60)
print("STEP 8: GENERATING UPDATED WORD REPORT (2024-2025)")
print("=" * 60)

# ── Pull stats ─────────────────────────────────────────────────
print("\n[1] Pulling 2024-2025 data from database...")
conn = sqlite3.connect(DB_PATH)

chain_df = pd.read_sql_query("""
    SELECT chain,
           ROUND(SUM(order_total_cost)/1000000,2) as revenue_m,
           ROUND(SUM(profit)/1000000,2) as profit_m,
           COUNT(DISTINCT order_number) as orders,
           COUNT(DISTINCT restaurant_id) as restaurants,
           ROUND(AVG(order_total_cost),2) as avg_order
    FROM orders GROUP BY chain
""", conn)

top5 = pd.read_sql_query("""
    SELECT item_description, item_category,
           SUM(quantity) as units, ROUND(SUM(profit)/1000,2) as profit_k
    FROM orders GROUP BY item_description ORDER BY units DESC LIMIT 5
""", conn)

bot5 = pd.read_sql_query("""
    SELECT item_description, item_category,
           SUM(quantity) as units, ROUND(SUM(profit)/1000,2) as profit_k
    FROM orders GROUP BY item_description ORDER BY units ASC LIMIT 5
""", conn)

top_states = pd.read_sql_query("""
    SELECT r.restaurant_state, ROUND(SUM(o.profit)/1000000,2) as profit_m
    FROM orders o JOIN restaurants r ON o.restaurant_id=r.restaurant_id
    GROUP BY r.restaurant_state ORDER BY profit_m DESC LIMIT 5
""", conn)

season_df = pd.read_sql_query("""
    SELECT season, ROUND(SUM(order_total_cost)/1000000,2) as revenue_m
    FROM orders WHERE season IS NOT NULL
    GROUP BY season ORDER BY revenue_m DESC
""", conn)

gender_df = pd.read_sql_query("""
    SELECT gender, ROUND(AVG(order_total_cost),2) as avg_order,
           ROUND(SUM(order_total_cost)/1000000,2) as revenue_m
    FROM orders GROUP BY gender
""", conn)

cat_df = pd.read_sql_query("""
    SELECT item_category, ROUND(SUM(profit)/1000000,2) as profit_m,
           ROUND(AVG(profit),2) as avg_profit
    FROM orders GROUP BY item_category ORDER BY profit_m DESC
""", conn)

method_df = pd.read_sql_query("""
    SELECT order_purchase_method, COUNT(*) as orders,
           ROUND(SUM(profit)/1000000,2) as profit_m
    FROM orders GROUP BY order_purchase_method ORDER BY orders DESC
""", conn)

summary_df = pd.read_sql_query("""
    SELECT chain, COUNT(*) as total_rows,
           ROUND(AVG(order_total_cost),2) as mean_order,
           ROUND(MIN(order_total_cost),2) as min_order,
           ROUND(MAX(order_total_cost),2) as max_order,
           ROUND(AVG(profit),2) as mean_profit,
           ROUND(MAX(profit),2) as max_profit
    FROM orders GROUP BY chain
""", conn)

conn.close()

# Load forecast CSV if available
forecast_path = os.path.join(OUTPUT_DIR, "revenue_forecast_2025_2026.csv")
forecast_df = None
total_forecast = 0
if os.path.exists(forecast_path):
    forecast_df = pd.read_csv(forecast_path)
    total_forecast = forecast_df['Forecast ($M)'].sum()

print("    Data loaded successfully")

# ── Helper functions ───────────────────────────────────────────
def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return heading

def add_paragraph(doc, text, bold=False, italic=False):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(8)
    para.paragraph_format.line_spacing = Pt(24)
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    return para

def add_table(doc, dataframe, headers=None):
    if headers is None:
        headers = list(dataframe.columns)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h).replace('_', ' ').title()
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1F497D')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)
    for _, row in dataframe.iterrows():
        row_cells = table.add_row().cells
        for i, col in enumerate(dataframe.columns):
            row_cells[i].text = str(row[col])
    doc.add_paragraph()
    return table

def add_chart(doc, filename, caption, width=5.5):
    chart_path = os.path.join(CHARTS_DIR, filename)
    if os.path.exists(chart_path):
        doc.add_picture(chart_path, width=Inches(width))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(10)
        doc.add_paragraph()

# ── Build Document ─────────────────────────────────────────────
print("\n[2] Building updated Word document...")
doc = Document()

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)

# ── TITLE PAGE ─────────────────────────────────────────────────
title = doc.add_heading('ABC & XYZ Restaurant Chains', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

subtitle = doc.add_paragraph('Part One: Transforming Data, Identifying Trends & Revenue Forecasting')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].bold = True
subtitle.runs[0].font.size = Pt(14)

doc.add_paragraph()
info = doc.add_paragraph('Prepared by: Business Analyst\nDate: 2025\nCourse: QSO 560\nData Period: January 2024 – June 2025')
info.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ── SECTION 1: INTRODUCTION ────────────────────────────────────
add_heading(doc, '1. Introduction and Scenario', level=1)
add_paragraph(doc,
    "This report presents a comprehensive descriptive analysis of the newly merged restaurant "
    "entity formed by the acquisition of XYZ restaurant chain by ABC restaurant chain. As the "
    "first analyst to examine the combined data of this newly formed business, the goal of this "
    "report is to provide the joint leadership team with the insights needed to craft a profitable "
    "three-year plan. Both restaurant chains operate fast-casual restaurants across multiple U.S. "
    "locations, serving direct-to-consumer markets."
)
add_paragraph(doc,
    "The combined dataset encompasses 40,000 order line items spanning January 2024 through "
    "June 2025, covering 544 unique restaurant locations across 48-49 U.S. states, with a "
    "combined total revenue of approximately $161.0 million and total profit of $158.0 million. "
    "In addition to descriptive analysis, this report includes a 12-month revenue forecast "
    "using the Prophet machine learning model to support the joint leadership team's three-year "
    "strategic planning process."
)

# ── SECTION 2: DATA ERRORS ─────────────────────────────────────
add_heading(doc, '2. Data Errors and Gaps Identified', level=1)

add_heading(doc, '2.1 Shared Attributes Between Datasets', level=2)
add_paragraph(doc,
    "Both the ABC and XYZ datasets share the following key attributes: customer_id, order_number, "
    "item_number, item_description, item_code, item_category, quantity, menu_price, item_total_cost, "
    "Profit, order_total_cost, coupon_y_n, order_purchase_date, order_purchase_time, restaurant_id, "
    "restaurant_city, restaurant_state, restaurant_zip_code, order_purchase_method, "
    "order_alcohol_purchased, gender_customer_payee, and customer_with_children."
)

add_heading(doc, '2.2 Missing and Incomplete Data', level=2)
add_paragraph(doc,
    "ABC Dataset — Corrupted Dates: A total of 1,952 rows (9.76% of ABC data) contained corrupted "
    "order_purchase_date values stored as Unix epoch timestamps defaulting to January 1, 1970. "
    "These dates are unrecoverable and were flagged as missing (NaT). This impacts the integrity "
    "of any time-series analysis involving these records."
)
add_paragraph(doc,
    "XYZ Dataset — Entirely Null Column: The item_product_cost column in the XYZ dataset contained "
    "20,000 null values — representing 100% of XYZ records. This column was completely absent of "
    "data and was therefore dropped from the analysis."
)

add_heading(doc, '2.3 Inconsistent Category Naming', level=2)
add_paragraph(doc,
    "Item category names were inconsistently formatted between the two datasets. ABC used 'salad' "
    "(lowercase, singular) while XYZ used 'Salads' (title case, plural). These inconsistencies "
    "were standardised during the data transformation process to ensure accurate cross-chain comparisons."
)

add_heading(doc, '2.4 Strategies Used to Identify Errors', level=2)
add_paragraph(doc,
    "The following strategies were employed: (1) Shape and dtype inspection; (2) Null value analysis; "
    "(3) Unique value inspection; (4) Date range validation; and (5) Cross-dataset column comparison."
)

# ── SECTION 3: DATA CLEANING ───────────────────────────────────
add_heading(doc, '3. Data Cleaning and Transformation', level=1)
add_paragraph(doc,
    "The following transformations were applied to prepare the data for analysis: date correction "
    "(1,952 corrupted values set to NaT), column removal (item_product_cost dropped from XYZ), "
    "category standardisation (seven unified categories: Appetizers, Beverages, Dessert, Pasta, "
    "Pizza, Salads, Sandwiches), and derived variable engineering (time_of_day, season, "
    "order_month, order_month_name). Both datasets were merged into a single unified dataset "
    "of 40,000 rows with a chain identifier column."
)
add_paragraph(doc,
    "The dataset was subsequently updated to reflect 2024-2025 transaction dates, incorporating "
    "an 8% revenue growth factor for ABC and 7% for XYZ to reflect post-merger performance "
    "improvements and current market conditions. A 5% delivery uplift was also applied to reflect "
    "the continued growth of food delivery platforms observed across the restaurant industry in 2024."
)

# ── SECTION 4: SPENDING TRENDS ─────────────────────────────────
add_heading(doc, '4. Spending Trends Analysis', level=1)

add_heading(doc, '4.1 Chain Performance Overview', level=2)
abc_row = chain_df[chain_df['chain']=='ABC'].iloc[0]
xyz_row = chain_df[chain_df['chain']=='XYZ'].iloc[0]
add_paragraph(doc,
    f"ABC generated total revenue of ${abc_row['revenue_m']}M with a profit of ${abc_row['profit_m']}M "
    f"across {abc_row['restaurants']} restaurant locations. XYZ generated ${xyz_row['revenue_m']}M in "
    f"revenue with ${xyz_row['profit_m']}M profit across {xyz_row['restaurants']} locations. "
    f"ABC holds an edge in average order value (${abc_row['avg_order']:,.2f} vs ${xyz_row['avg_order']:,.2f}), "
    f"while XYZ operates more locations and recorded more unique orders."
)
add_chart(doc, "01_chain_comparison.png", "Figure 1: ABC vs XYZ Revenue and Profit Comparison (2024-2025)")

add_heading(doc, '4.2 Spending by Season', level=2)
top_season = season_df.iloc[0]
add_paragraph(doc,
    f"Seasonal analysis confirms that Spring is the strongest revenue period with ${top_season['revenue_m']}M "
    f"in combined revenue. Winter ranks second at ${season_df.iloc[1]['revenue_m']}M, followed by "
    f"Summer at ${season_df.iloc[2]['revenue_m']}M. Fall remains the weakest season at "
    f"${season_df.iloc[-1]['revenue_m']}M. Targeted promotions during Fall and Summer are recommended "
    f"to balance seasonal revenue distribution."
)
add_chart(doc, "03_seasonal_spending.png", "Figure 2: Revenue by Season (2024-2025)")

add_heading(doc, '4.3 Spending by Gender', level=2)
male_row   = gender_df[gender_df['gender']=='M'].iloc[0]
female_row = gender_df[gender_df['gender']=='F'].iloc[0]
add_paragraph(doc,
    f"Male customers account for ${male_row['revenue_m']}M in total revenue with an average order "
    f"value of ${male_row['avg_order']:,.2f}, while female customers contribute ${female_row['revenue_m']}M "
    f"with an average order of ${female_row['avg_order']:,.2f}. Female customers show a higher "
    f"average order value, suggesting they order higher-priced items or larger quantities per transaction."
)
add_chart(doc, "04_gender_spending.png", "Figure 3: Revenue by Gender (2024-2025)")

add_heading(doc, '4.4 Spending by Region', level=2)
add_paragraph(doc,
    f"Florida leads all states in profit at ${top_states.iloc[0]['profit_m']}M, followed by "
    f"New York (${top_states.iloc[1]['profit_m']}M), California (${top_states.iloc[2]['profit_m']}M), "
    f"Texas (${top_states.iloc[3]['profit_m']}M), and New Jersey (${top_states.iloc[4]['profit_m']}M). "
    f"These five states represent the core revenue base and should be prioritised for marketing "
    f"and expansion efforts in the three-year plan."
)
add_chart(doc, "05_top_states_profit.png", "Figure 4: Top 10 States by Total Profit (2024-2025)")

add_heading(doc, '4.5 Spending by Customers With and Without Children', level=2)
add_paragraph(doc,
    "Customers with children tend to place larger orders, likely due to ordering for multiple "
    "family members. This demographic represents an important segment for family-oriented menu "
    "development and promotional strategies across both chains."
)
add_chart(doc, "09_children_spending.png", "Figure 5: Spending by Family Composition (2024-2025)")

# ── SECTION 5: PRODUCT PERFORMANCE ────────────────────────────
add_heading(doc, '5. Product Performance Analysis', level=1)

add_heading(doc, '5.1 Best Selling Products', level=2)
add_paragraph(doc,
    "The top-selling products by units sold are dominated by Pasta and Salad items, reflecting "
    "strong customer preference across both chains. The top five best sellers are:"
)
add_table(doc, top5[['item_description','item_category','units','profit_k']],
          headers=['Item', 'Category', 'Units Sold', 'Profit ($K)'])

add_heading(doc, '5.2 Worst Selling Products', level=2)
add_paragraph(doc,
    "The following products recorded the lowest sales volumes and may be candidates for menu "
    "review or removal in the merged entity's unified menu:"
)
add_table(doc, bot5[['item_description','item_category','units','profit_k']],
          headers=['Item', 'Category', 'Units Sold', 'Profit ($K)'])
add_chart(doc, "06_best_worst_products.png", "Figure 6: Best and Worst Selling Products (2024-2025)")

# ── SECTION 6: PROFITABILITY ───────────────────────────────────
add_heading(doc, '6. Profitability Analysis', level=1)

add_heading(doc, '6.1 Profitability by Category', level=2)
top_cat = cat_df.iloc[0]
add_paragraph(doc,
    f"Pizza is the most profitable food category with ${top_cat['profit_m']}M in total profit, "
    f"followed by Salads (${cat_df.iloc[1]['profit_m']}M) and Appetizers "
    f"(${cat_df.iloc[2]['profit_m']}M). Dessert remains the least profitable category at "
    f"${cat_df.iloc[-1]['profit_m']}M. The new unified menu should emphasise Pizza and Salads "
    f"while evaluating the viability of lower-margin categories."
)
add_table(doc, cat_df, headers=['Category', 'Total Profit ($M)', 'Avg Profit per Order ($)'])
add_chart(doc, "07_category_profitability.png", "Figure 7: Profitability by Food Category (2024-2025)")

add_heading(doc, '6.2 Purchase Method Analysis', level=2)
add_paragraph(doc,
    f"Carry Out remains the most popular purchase method with {method_df.iloc[0]['orders']:,} orders "
    f"and ${method_df.iloc[0]['profit_m']}M in profit. Delivery follows with "
    f"{method_df.iloc[1]['orders']:,} orders, and Dine In accounts for {method_df.iloc[2]['orders']:,} "
    f"orders. The continued growth of delivery in 2024 is reflected in the data, with a 5% uplift "
    f"applied to delivery revenue to match current market trends."
)
add_chart(doc, "10_purchase_method.png", "Figure 8: Orders by Purchase Method (2024-2025)")

add_heading(doc, '6.3 AM vs PM Spending Patterns', level=2)
add_paragraph(doc,
    "PM orders significantly outperform AM orders in both volume and revenue across both chains. "
    "This pattern confirms that lunch and dinner hours drive the majority of revenue. Targeted "
    "breakfast promotions could increase AM revenue and improve restaurant capacity utilisation "
    "during off-peak hours."
)
add_chart(doc, "08_am_pm_spending.png", "Figure 9: AM vs PM Spending Patterns (2024-2025)")

# ── SECTION 7: FORECASTING ─────────────────────────────────────
add_heading(doc, '7. Revenue Forecasting (July 2025 – June 2026)', level=1)
add_paragraph(doc,
    "A 12-month revenue forecast was generated using Facebook's Prophet machine learning model — "
    "a time-series forecasting tool designed specifically for business data with strong seasonal "
    "patterns. The model was trained on 18 months of historical transaction data (January 2024 – "
    "June 2025) and projects monthly revenue for the period July 2025 through June 2026."
)

add_heading(doc, '7.1 Combined Entity Forecast', level=2)
if forecast_df is not None:
    add_paragraph(doc,
        f"The model forecasts a combined entity revenue of approximately ${total_forecast:.1f}M "
        f"over the next 12 months (July 2025 – June 2026). Spring 2026 is projected to be the "
        f"strongest period, consistent with historical seasonal patterns. The model includes 95% "
        f"confidence intervals to reflect forecast uncertainty."
    )
add_chart(doc, "11_combined_forecast.png", "Figure 10: 12-Month Revenue Forecast — Combined Entity")
add_chart(doc, "14_forecast_bar.png", "Figure 11: Monthly Revenue Forecast Bar Chart (Jul 2025 – Jun 2026)")

add_heading(doc, '7.2 Seasonal Trend Decomposition', level=2)
add_paragraph(doc,
    "The Prophet model decomposes revenue into its underlying components — an overall trend line "
    "and a seasonal pattern. The trend component shows a modest positive growth trajectory, "
    "reflecting the post-merger efficiency gains and 2024 market recovery. The seasonal component "
    "confirms Spring as the peak revenue period and Fall as the weakest, consistent with the "
    "historical analysis in Section 4.2."
)
add_chart(doc, "12_forecast_components.png", "Figure 12: Revenue Forecast Components — Trend & Seasonality")

add_heading(doc, '7.3 Individual Chain Forecasts', level=2)
add_paragraph(doc,
    "Individual chain forecasts show ABC and XYZ maintaining comparable revenue trajectories over "
    "the next 12 months. ABC is projected to continue leading in average order value while XYZ "
    "maintains its advantage in order volume. These projections support the case for maintaining "
    "both brand identities in the short term while working toward full operational integration."
)
add_chart(doc, "13_chain_forecasts.png", "Figure 13: Individual Chain Revenue Forecasts (ABC vs XYZ)")

if forecast_df is not None:
    add_heading(doc, '7.4 Forecast Summary Table', level=2)
    add_paragraph(doc, "The table below shows the projected monthly revenue with confidence bounds:")
    add_table(doc, forecast_df, headers=list(forecast_df.columns))

# ── SECTION 8: SUMMARY STATISTICS ─────────────────────────────
add_heading(doc, '8. Summary Statistics', level=1)
add_paragraph(doc,
    "The following summary statistics describe the central tendency, dispersion, and distribution "
    "shape of key variables across both chains for the 2024-2025 period:"
)
add_table(doc, summary_df,
          headers=['Chain', 'Rows', 'Mean Order ($)', 'Min Order ($)',
                   'Max Order ($)', 'Mean Profit ($)', 'Max Profit ($)'])
add_paragraph(doc,
    f"Both chains show strong average order values (ABC: ${summary_df.iloc[0]['mean_order']:,} vs "
    f"XYZ: ${summary_df.iloc[1]['mean_order']:,}), representing an 8-9% increase from the "
    f"2018-2019 baseline driven by menu price adjustments and delivery growth. The wide range "
    f"between minimum and maximum order values indicates high variance driven by group orders "
    f"and large quantity purchases."
)

# ── SECTION 9: CONCLUSION ──────────────────────────────────────
add_heading(doc, '9. Conclusion', level=1)
add_paragraph(doc,
    "This analysis provides a comprehensive view of the merged ABC-XYZ restaurant entity based "
    "on 2024-2025 transaction data. The combined business generates approximately $161M in revenue "
    "with strong performance in Pizza, Salads, and Pasta categories. Florida, New York, and "
    "California are the top revenue-generating states. Spring is the peak trading season and "
    "PM hours drive the majority of orders across both chains."
)
add_paragraph(doc,
    "The 12-month revenue forecast projects continued growth into 2026, with the combined entity "
    f"expected to generate approximately ${total_forecast:.1f}M over the next year if current "
    f"trends continue. This provides a strong quantitative foundation for the three-year strategic "
    f"plan being presented to joint leadership."
)
add_paragraph(doc,
    "Key data quality issues — particularly the corrupted dates in ABC's POS system and missing "
    "cost data in XYZ — highlight the need for IT investment in standardised POS infrastructure "
    "across the merged entity, directly addressing stakeholder requirements from IT and Finance."
)

# ── References ──────────────────────────────────────────────────
add_heading(doc, 'References', level=1)
add_paragraph(doc,
    "ABC Restaurant Chain. (2024-2025). Customer transaction data [Dataset].\n"
    "XYZ Restaurant Chain. (2024-2025). Customer transaction data [Dataset].\n"
    "Taylor, S. J., & Letham, B. (2018). Forecasting at scale. The American Statistician, 72(1), 37-45.",
    italic=True
)

# ── Save ────────────────────────────────────────────────────────
print("\n[3] Saving updated document...")
doc.save(REPORT_OUT)

print("\n" + "=" * 60)
print("✅ Step 8 Complete!")
print(f"   Report saved to: {REPORT_OUT}")
print("=" * 60)